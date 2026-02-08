import docx

def replace_following_text(doc, heading_text, new_content):
    paragraphs = doc.paragraphs
    for i, p in enumerate(paragraphs):
        if heading_text in p.text:
            # Look at the next paragraph
            if i + 1 < len(paragraphs):
                target_p = paragraphs[i + 1]
                # Modifying runs to try and preserve style
                if target_p.runs:
                    target_p.runs[0].text = new_content
                    # Clear remaining runs to avoid duplicate/messy text
                    for run in target_p.runs[1:]:
                        run.text = ''
                else:
                    target_p.text = new_content
                print(f"Updated section: {heading_text}")
            return
    print(f"WARNING: Section not found: {heading_text}")

# Detailed Content Dictionary
content_map = {
    "Problem Statement:": 
"""Memory safety vulnerabilities, specifically buffer overflows and pointer corruption, remain one of the most critical classes of security flaws in modern software. Standard defenses like Address Space Layout Randomization (ASLR) provide a significant barrier but are often bypassed through information leaks (e.g., format string vulnerabilities or side-channels). Once a single memory address is leaked, the randomization is defeated, allowing unrestricted access to the process memory. Furthermore, traditional page-level protections (mprotect) are often too coarse and expensive for frequent in-process context switching, leaving a gap in defense-in-depth strategies.""",

    "Background Information: (literature review)":
"""The 'arms race' between attackers and defenders has evolved from simple stack smashing protection (Stack Cookies) to randomization (ASLR) and non-executable memory (DEP/NX). However, advanced exploit techniques like Return-Oriented Programming (ROP) and Just-In-Time (JIT) spraying have continuously challenged these defenses. Intel's Memory Protection Keys (MPK/PKU) offer a hardware-based solution for thread-local permission switching, allowing for performant compartmentalization without the overhead of context switches. This project builds upon these concepts, simulating a hybrid architecture that combines the layout randomization of ASLR with the strict access control of MPK.""",

    "Primary Objectives:":
"""1. To develop a vulnerable C application demonstrating classic stack-based buffer overflow and pointer hijacking mechanisms.\n2. To implement an ASLR-enabled scenario and demonstrate how standard ASLR is bypassed via information leaks.\n3. To implement a Hybrid Protection System (ASLR + MPK Simulation) that utilizes distinct memory zones (Buffer vs. Secret) and enforces access control, preventing data theft even when the attacker knows the memory layout.""",

    "Secondary Objectives:":
"""1. To create a real-time Interactive GUI Dashboard that visualizes memory operations, attack vectors, and protection mechanisms for educational purposes.\n2. To provide a comparative analysis of attack success rates across Baseline, ASLR, and Protected scenarios.\n3. To demonstrate the 'Fail-Closed' principle where hardware traps block unauthorized access attempts.""",

    "4.1 Approach:":
"""The project adopts a 'Red Team vs. Blue Team' simulation approach. The Red Team component involves a custom Python-based exploit generator that calculates offsets, parses leaked addresses from stdout, and constructs malicious payloads. The Blue Team component involves a C application designated with specific 'Zones' (Zone U for User/Buffer, Zone S for Secret). \n\nConceptual Flow:\n[Attacker] -> [Inject Payload] -> [Buffer Overflow] -> [Overwrite Pointer] -> [Access Secret] ... [Protection Check] -> [Block/Trap].""",

    "4.2 Procedures:":
"""1. Baseline Implementation: Developed 'vulnerable_app.c' with no protection to establish a control case.\n2. ASLR Integration: Modified the app to allocate memory at random offsets ('aslr_app.c').\n3. Protection Logic: Implemented 'protected_app.c' using a monitor pattern to simulate MPK checks.\n4. Exploit Development: Wrote 'run_exploit.py' to automate the attack chain (Leak -> Calculate -> Pwn).\n5. GUI Development: Built a Flask-based frontend with EventSource streaming to visualize the binary memory slots in real-time.""",

    "5.1 Planning and Design:":
"""The memory model was designed with a custom struct 'MemoryLayout' to facilitate easier visualization of the 'Buffer Overflow to Pointer Hijack' path. We specifically chose to implement a user-space simulation of MPK logic to ensure the demonstration could be run portability on non-server hardware while still proving the architectural concepts of fine-grained compartmentalization.""",

    "5.2 Implementation:":
"""The backend was implemented in C using GCC, defining 'Zone' structures where the Protected app separates the Secret execution domain from the Buffer domain. The Python interfacing layer uses subprocess.Popen for real-time interaction with the compiled binaries, capturing stdout to drive the GUI. The Cyberpunk-themed GUI was built using HTML5/CSS3 to make abstract memory concepts visually engaging.""",

    "6.1 Tools:":
"""- GCC: Compiler for C target applications.\n- Python 3: Exploit engine and Flask web server.\n- Flask: Web framework for dashboard.\n- HTML5/CSS3/JavaScript: Frontend visualization.\n- Bash: Orchestration scripts.\n- GDB: Used during development for debugging memory offsets.""",

    "6.2 Techniques:":
"""- Input Stream Injection: Sending raw bytes to stdin to simulate network payloads.\n- Address Leak Parsing: Regex matching on process output to simulate reading leaked pointers.\n- Server-Sent Events (SSE): For low-latency streaming of terminal logs.\n- DOM Animation: CSS3 animations to visualize buffer filling and protection shields.""",

    "7.1 Initial Findings:":
"""In the Baseline scenario, the secret was compromised 100% of the time, with the dashboard correctly visualizing the overflow. In the ASLR scenario, the attack crashed (Segfault) without leaks, but succeeded immediately once an address leak was introduced, confirming ASLR's brittleness against leaks.""",

    "7.2 Iterative Improvements:":
"""Initial animations were found to be too 'game-like' (using projectiles). These were refined into professional 'Data Flow' and 'Signal Integrity' visualizations. Additionally, a 'Hidden by Default' privacy mode was implemented for the secret data to better represent secure memory contents.""",

    "8.1 Final Results:":
"""- Scenario 1 (Unprotected): Attack Success. Secret overwritten.\n- Scenario 2 (ASLR Only): Attack Success (with leak). Randomization bypassed.\n- Scenario 3 (Hybrid Protection): Attack Blocked. Even with the correct address, the dereference triggered the simulated hardware trap (Access Denied).""",

    "8.2 Discussion:":
"""The results validate the hypothesis: Architecture > Obscurity. ASLR relies on the secrecy of addresses. Once lost, the defense collapses. The Hybrid MPK approach relies on Authorization. Even if the attacker knows 'where' the data is, they lack the 'key' to access it, demonstrating a Zero-Trust architecture applied to memory.""",

    "9.1 Prototype Description:":
"""The prototype is a self-contained 'Memory Defense Dashboard'. It features an Interactive Console for scenario selection, a dynamic Memory Visualizer grid representing RAM bytes, pointers, and secret zones, and a Real-Time Terminal showing the attacker's shell output.""",

    "9.2 Development Process:":
"""Development proceeded in three phases: 1. Core C Logic & Exploits, 2. Basic Web Interface & Connectivity, 3. Advanced Animations & Polish. This modular approach allowed for independent testing of the exploit logic before visual integration.""",

    "9.3 Testing and Validation:":
"""Validated against generic buffer overflow payloads and verified 'Fail-Closed' behavior. Cross-verified across multiple runs to ensure ASLR randomization was effective (addresses changed every run).""",

    "10.1 Summary:":
"""This project successfully demonstrated the limitations of current ASLR implementations and the robustness of fine-grained compartmentalization. We built a fully functional end-to-end demonstration platform that makes complex memory security concepts accessible and visible.""",

    "10.2 Personal Reflection:":
"""Working on this project highlighted the fragility of low-level software. Seeing how easily a single missing bound check can compromise an entire system—and conversely, how a strong architectural guarantee (MPK) can mitigate it—was a profound learning experience regarding 'Secure by Design' principles.""",

    "11. Visuals:":
"""(See connected Dashboard for live visuals)\n1. Buffer Overflow: Red byte wave consuming safe buffer.\n2. Pointer Corruption: Pointer value glitching to target address.\n3. Protection Shield: Lock icon pulsing and rejecting access.""",

    "12. Outcome of the work:":
"""- Product: Deployable 'Memory Defense Dashboard' tool.\n- Codebase: Open-source GitHub repository 'Memory-Compartmentalization-using-MPK-keys'.\n- Documentation: Comprehensive walkthroughs and technical reports."""
}

doc = docx.Document("4. EL report.docx")

# Sort keys by length descending to match longest headers first if overlaps exist
sorted_keys = sorted(content_map.keys(), key=len, reverse=True)

for i, p in enumerate(doc.paragraphs):
    p_text = p.text.strip()
    match_key = None
    
    # specialized fix for 'Visuals:' which might be '11. Visuals:' in doc
    # normalized check
    for key in sorted_keys:
        # Check for strict start match (e.g. "Problem Statement:")
        if p_text.lower().startswith(key.lower()) or (key in p_text and len(p_text) < len(key) + 5):
            match_key = key
            break
            
    if match_key:
        content = content_map[match_key]
        
        # Determine if it's a Header-Only line or Inline
        # If the paragraph is significantly longer than the key, it's likely Inline (Header + Placeholder)
        # Exception: "Introduction" (len 12) vs "Probability..."
        
        # Case 1: Header Only (e.g., "Problem Statement:")
        # We assume the user creates content in the NEXT paragraph.
        # Threshold: if paragraph is roughly the same length as key (allow some whitespace/typos)
        if len(p_text) <= len(match_key) + 5: 
            # Replace NEXT paragraph
            if i + 1 < len(doc.paragraphs):
                target_p = doc.paragraphs[i+1]
                # Avoid overwriting if next paragraph is ALSO a header
                # (Simple heuristic: don't overwrite if it matches another key)
                is_next_header = False
                for k in sorted_keys:
                    if target_p.text.strip().lower().startswith(k.lower()):
                        is_next_header = True
                        break
                
                if not is_next_header:
                    if target_p.runs:
                        target_p.runs[0].text = content
                        for run in target_p.runs[1:]:
                            run.text = ''
                    else:
                        target_p.text = content
                    print(f"Updated Section (Next Line): {match_key}")
                else:
                    print(f"SKIPPED Overwrite for {match_key} because next line looks like a header: {target_p.text[:20]}")
        
        # Case 2: Inline (e.g., "8.2 Discussion: Interpret the results...")
        else:
            # It's inline. We replace the current paragraph text but keep the header.
            # Using simple text replacement for reliability
            p.text = f"{match_key} {content}"
            print(f"Updated Section (Inline): {match_key}")

output_filename = "Project_Report_Final.docx"
doc.save(output_filename)
print(f"Successfully generated {output_filename}")
